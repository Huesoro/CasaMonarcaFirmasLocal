import os
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import shutil
import logging
from pyhanko.sign import signers
from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
from Azure_SQL.conectBases import firmar_documento

logging.basicConfig(level=logging.INFO, filename="debug_firmas.log")

def buscar_usuario_por_id(user_id):
    import sqlite3
    conn = sqlite3.connect('Azure_SQL/Usuarios.db')
    cursor = conn.cursor()
    cursor.execute("SELECT user_id, name, role, email FROM usuarios WHERE user_id = ?", (user_id,))
    resultado = cursor.fetchone()
    conn.close()
    if resultado:
        return {'user_id': resultado[0], 'name': resultado[1], 'role': resultado[2], 'email': resultado[3]}
    else:
        return None

def buscar_documento_por_id(doc_id):
    import sqlite3
    conn = sqlite3.connect('Azure_SQL/Documentos.db')
    cursor = conn.cursor()
    cursor.execute("SELECT doc_id, title, sharepoint_url FROM documentos WHERE doc_id = ?", (doc_id,))
    resultado = cursor.fetchone()
    conn.close()
    if resultado:
        return {'doc_id': resultado[0], 'title': resultado[1], 'sharepoint_url': resultado[2]}
    else:
        return None

def buscar_archivo_sharepoint(sharepoint_url):
    ruta = sharepoint_url
    if os.path.exists(ruta):
        return ruta
    else:
        raise FileNotFoundError(f"Archivo no encontrado en SharePoint: {ruta}")

def normalizar_correo(correo):
    return correo.strip().lower().replace("@", "_at_").replace(".", "_dot_")

def buscar_firma_por_correo(correo):
    archivo = normalizar_correo(correo) + ".png"
    ruta = os.path.join('SharePoint/firmas', archivo)
    if os.path.exists(ruta):
        return ruta
    else:
        raise FileNotFoundError(f"Firma no encontrada para {correo}")

def firmar_documento_word(ruta_docx, ruta_firma, nombre, rol):
    doc = Document(ruta_docx)
    firmado = False
    rol_map = {
        "reception": {"marcador": "{nombre_receptor}", "texto_rol": "(Personal autorizado)"},
        "finance": {"marcador": "{nombre_finanzas}", "texto_rol": "(responsable de finanzas)"},
        "admin": {"marcador": "{nombre_dueno}", "texto_rol": "(Representante legal de Casa Monarca)"},
        "inventory": {"marcador": "{nombre_inventario}", "texto_rol": "(Responsable de inventarios)"}
    }
    if rol not in rol_map:
        print(f"Rol {rol} no reconocido para firma en Word.")
        return ruta_docx
    marcador = rol_map[rol]["marcador"]
    texto_rol = rol_map[rol]["texto_rol"]
    print(f"Buscando marcador: {marcador} para rol: {rol}")
    print(f"Tablas en el documento: {len(doc.tables)}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for i, paragraph in enumerate(cell.paragraphs):
                    if marcador in paragraph.text:
                        print(f"Marcador encontrado en tabla: {paragraph.text}")
                        if not firmado:
                            parrafo_firma = cell.paragraphs[i].insert_paragraph_before()
                            run = parrafo_firma.add_run()
                            run.add_picture(ruta_firma, width=Inches(2))
                            parrafo_linea = paragraph.insert_paragraph_before()
                            parrafo_linea.add_run("____________________")
                        nuevo_texto = paragraph.text.replace(marcador, nombre)
                        paragraph.text = nuevo_texto
                        firmado = True
    if firmado:
        doc.save(ruta_docx)
        print(f"Documento firmado y guardado en: {ruta_docx}")
        logging.info("Documento firmado correctamente.")
    else:
        print("No se encontró el marcador para firmar en el documento.")
        logging.info("No se encontró el marcador para firmar en el documento.")
    return ruta_docx

def buscar_certificado_correo(correo):
    archivo = correo+"_cert.pem"
    ruta = os.path.join('certs', archivo)
    if os.path.exists(ruta):
        return ruta
    else:
        raise FileNotFoundError(f"Certificado no encontrado para {correo}")

def buscar_llave_correo(correo):
    archivo = correo+"_key.pem"
    ruta = os.path.join('certs', archivo)
    if os.path.exists(ruta):
        return ruta
    else:
        raise FileNotFoundError(f"Llave no encontrada para {correo}")

async def firmar_digitalmente_pdf(pdf_path, cert_path, key_path, output_path, password):
    import shutil
    print("DEBUG cert_path:", cert_path, "type:", type(cert_path))
    print("DEBUG key_path:", key_path, "type:", type(key_path))
    print("DEBUG existe cert_path:", os.path.exists(cert_path))
    print("DEBUG existe key_path:", os.path.exists(key_path))
    cms_signer = signers.SimpleSigner.load(
        key_path,
        cert_path,
        key_passphrase=password.encode()
    )
    with open(pdf_path, 'rb') as doc:
        w = IncrementalPdfFileWriter(doc, strict=False)
        out = await signers.async_sign_pdf(
            w,
            signers.PdfSignatureMetadata(field_name='FirmaDigital', reason='Firma digital Casa Monarca'),
            signer=cms_signer,
        )
        with open(output_path, 'wb') as outf:
            shutil.copyfileobj(out, outf)

async def flujo_documento(user_id, doc_id, password_firma, placeholder_nombre="{nombre_inventario}", rol=None):
    try:
        usuario = buscar_usuario_por_id(user_id)
        logging.info(f"Usuario encontrado: {usuario}")
        if not usuario:
            logging.info("Usuario no encontrado")
            return {"status": "error", "message": "Usuario no encontrado"}
        documento = buscar_documento_por_id(doc_id)
        logging.info(f"Documento encontrado: {documento}")
        if not documento:
            logging.info("Documento no encontrado")
            return {"status": "error", "message": "Documento no encontrado"}
        try:
            ruta_archivo = buscar_archivo_sharepoint(documento['sharepoint_url'])
            logging.info(f"Archivo encontrado: {ruta_archivo}")
        except FileNotFoundError as e:
            logging.info(f"Archivo no encontrado: {e}")
            return {"status": "error", "message": str(e)}
        if not os.path.isfile(ruta_archivo):
            logging.info(f"El archivo no existe o no es accesible: {ruta_archivo}")
            return {"status": "error", "message": f"El archivo no existe: {ruta_archivo}"}
            
        carpeta_destino = os.path.dirname(ruta_archivo)
        nombre_archivo = os.path.basename(ruta_archivo)
        
        # Crear copias temporales para el proceso de firma
        temp_docx = os.path.join(carpeta_destino, f"temp_{nombre_archivo}")
        temp_pdf = os.path.join(carpeta_destino, f"temp_{nombre_archivo.replace('.docx', '.pdf')}")
        
        # Copiar el archivo original a un temporal para trabajar
        shutil.copy2(ruta_archivo, temp_docx)
        
        try:
            ruta_firma = buscar_firma_por_correo(usuario['email'])
        except FileNotFoundError as e:
            logging.info(f"Firma no encontrada: {e}")
            return {"status": "error", "message": str(e)}
        try:
            ruta_certificado = buscar_certificado_correo(usuario['email'])
        except FileNotFoundError as e:
            logging.info(f"Certificado no encontrado: {e}")
            return {"status": "error", "message": str(e)}
        try:
            ruta_llave = buscar_llave_correo(usuario['email'])
        except FileNotFoundError as e:
            logging.info(f"Llave no encontrada: {e}")
            return {"status": "error", "message": str(e)}
            
        rol_firma = rol if rol else usuario['role'] if usuario['role'] else "Personal autorizado"
        
        # Firmar el documento Word temporal
        firmar_documento_word(temp_docx, ruta_firma, usuario['name'], rol_firma)
        
        try:
            # Convertir a PDF temporal
            convert(temp_docx, temp_pdf)
            logging.info(f"Convertido a PDF: {temp_pdf}")
        except Exception as e:
            logging.info(f"Error al convertir a PDF: {e}")
            os.remove(temp_docx)  # Limpiar archivo temporal
            return {"status": "error", "message": f"Error al convertir a PDF: {e}"}
            
        try:
            # Firmar digitalmente el PDF temporal
            await firmar_digitalmente_pdf(temp_pdf, ruta_certificado, ruta_llave, ruta_archivo.replace('.docx', '.pdf'), password_firma)
            logging.info(f"PDF firmado digitalmente")
        except Exception as e:
            logging.info(f"Error al firmar digitalmente: {e}")
            os.remove(temp_docx)  # Limpiar archivos temporales
            os.remove(temp_pdf)
            return {"status": "error", "message": f"Error al firmar digitalmente: {e}"}
            
        try:
            # Sobrescribir el archivo original con el temporal firmado
            shutil.move(temp_docx, ruta_archivo)
            
            # Actualizar el estado en la base de datos
            firmar_documento(doc_id, user_id, rol_firma)
            logging.info(f"Flujo de firmas actualizado para doc_id={doc_id}, user_id={user_id}, rol={rol_firma}")
        except Exception as e:
            logging.info(f"Error al actualizar flujo de firmas: {e}")
            return {"status": "error", "message": f"Error al actualizar flujo de firmas: {e}"}
            
        # Limpiar archivos temporales si quedan
        try:
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
        except:
            pass
            
        return {
            "status": "success",
            "message": "Documento firmado y guardado correctamente.",
            "documento": documento
        }
    except Exception as e:
        logging.info(f"ERROR FLUJO: {e}")
        # Intentar limpiar archivos temporales en caso de error
        try:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
        except:
            pass
        return {"status": "error", "message": f"Error inesperado en el flujo: {e}"}
