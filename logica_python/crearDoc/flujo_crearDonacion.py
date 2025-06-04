import os
import uuid
from datetime import datetime
from Azure_SQL.conectBases import insert_doc, insert_donacion_dinero, insert_donacion_insumos, insertar_flujo_firmas
from docx import Document


DOCUMENTOS_PATH = "SharePoint/Donaciones/"
MACHOTES_PATH = "SharePoint/Machotes/"

def rellenar_campos_docx(path_machote, datos, id_documento):
    """
    Rellena los campos {campo} en un documento .docx y guarda usando un ID único.

    Args:
        path_machote (str): Ruta del documento machote (.docx) a usar.
        datos (dict): Diccionario con claves {campo} y sus valores correspondientes.
        id_documento (str or int): ID único del documento (ejemplo: 123456)
    """
    # Cargar el documento
    doc = Document(path_machote)

    # Reemplazar en cada párrafo
    for p in doc.paragraphs:
        for key, value in datos.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))

    # Reemplazar dentro de tablas también
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in datos.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Definir ruta de salida basada en el ID
    path_salida = os.path.join(DOCUMENTOS_PATH, f"doc_{id_documento}.docx")

    # Guardar el documento rellenado en su nueva ruta
    doc.save(path_salida)
    return path_salida

def crear_donacion(datos, user_id):
    """
    Crea el documento de donación, lo guarda y registra la donación en la base de datos.
    datos: dict con los datos a rellenar en el docx.
    user_id: id del usuario que dona.
    """
    # 1. Generar ID único para el documento
    id_documento = str(uuid.uuid4())
    now = datetime.now()
    fecha_hoy = now.strftime("%Y-%m-%d %H:%M:%S")

    # 2. Determinar tipo de donación y machote
    if "{monto_donado}" in datos or "monto_donado" in datos:
        tipo_donacion = "dinero"
        path_machote = os.path.join(MACHOTES_PATH, "Donacion_Dinero.docx")
    elif "{lista_articulos}" in datos or "lista_articulos" in datos:
        tipo_donacion = "especie"
        path_machote = os.path.join(MACHOTES_PATH, "Donacion_Especie.docx")
    else:
        return {"status": "error", "message": "No se reconoce el tipo de donación."}

    # 3. Rellenar y guardar el documento
    path_docx = rellenar_campos_docx(path_machote, datos, id_documento)
    
    # Obtener el nombre del donante
    nombre_donante = datos.get("nombre_donante") or datos.get("{nombre_donante}", "Desconocido")

    # 4. Registrar el documento en la base de datos
    doc_data = {
        "doc_id": id_documento,
        "title": f"Donación {tipo_donacion} {nombre_donante}",
        "sharepoint_url": path_docx,
        "Type": tipo_donacion,
        "status": "pending",
        "user_id": user_id,
        "created_at": fecha_hoy,
        "updated_at": fecha_hoy
    }
    insert_doc(doc_data)

    # 5. Registrar la donación según el tipo
    if tipo_donacion == "dinero":
        don_data = {
            "amount": datos.get("monto_donado") or datos.get("{monto_donado}"),
            "user_id": user_id,
            "created_at": fecha_hoy,
            "updated_at": fecha_hoy,
            "doc_id": id_documento
        }
        insert_donacion_dinero(don_data)
    elif tipo_donacion == "especie":
        insumo_data = {
            "amount": datos.get("cantidad") or 1,  # Puedes ajustar esto según tu lógica
            "objeto": datos.get("lista_articulos") or datos.get("{lista_articulos}"),
            "user_id": user_id,
            "created_at": fecha_hoy,
            "updated_at": fecha_hoy,
            "doc_id": id_documento
        }
        insert_donacion_insumos(insumo_data)

    # 6. Insertar el flujo de firmas
    insertar_flujo_firmas(id_documento, user_id, tipo_donacion)

    return {
        "status": "success",
        "message": "Donación registrada correctamente.",
        "tipo_donacion": tipo_donacion,
        "id_documento": id_documento,
        "path_docx": path_docx
    }
